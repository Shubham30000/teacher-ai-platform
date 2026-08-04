"""
DOCX parser built on ``python-docx``.

Walks the document body in document order (not ``document.paragraphs``
and ``document.tables`` separately, which loses interleaving) so that
headings, paragraphs, and tables end up nested in the same order they
appear in the source file. Images are extracted from inline shapes and
saved alongside the upload; hyperlinks are extracted from paragraph
runs.
"""
from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Optional

from app.core.constants import ElementType, SupportedFileType
from app.core.exceptions import ParsingError
from app.document_intelligence.models import (
    ContentBlock,
    DocumentMetadata,
    HeadingLevel,
    ImageElement,
    Section,
    StructuredDocument,
    TableCell,
    TableElement,
    UrlElement,
)
from app.parsers.base import BaseParser

logger = logging.getLogger(__name__)

_URL_RE = re.compile(r"https?://[^\s)\]}'\"<>]+")
_HEADING_STYLE_RE = re.compile(r"heading\s*(\d)", re.IGNORECASE)


def _heading_level_from_style(style_name: str) -> Optional[HeadingLevel]:
    if not style_name:
        return None
    if style_name.strip().lower() in ("title",):
        return HeadingLevel.TITLE
    match = _HEADING_STYLE_RE.search(style_name)
    if match:
        level = min(int(match.group(1)), HeadingLevel.H4.value)
        return HeadingLevel(level)
    return None


class DocxParser(BaseParser):
    format_name = "docx"
    supports_ocr = False

    def parse(self, file_path: Path) -> StructuredDocument:
        self._base_check(file_path)
        try:
            import docx
            from docx.document import Document as DocxDocument
            from docx.oxml.table import CT_Tbl
            from docx.oxml.text.paragraph import CT_P
            from docx.table import Table
            from docx.text.paragraph import Paragraph
        except ImportError as exc:
            raise ParsingError(
                "python-docx is not installed. Install it with `pip install python-docx`."
            ) from exc

        try:
            document: "DocxDocument" = docx.Document(str(file_path))
        except Exception as exc:  # noqa: BLE001 - surface as domain error
            raise ParsingError(f"Failed to open DOCX file: {exc}") from exc

        root_sections: list[Section] = []
        stack: list[tuple[int, Section]] = []
        preamble = Section(heading=None, heading_level=HeadingLevel.TITLE)
        root_sections.append(preamble)
        current: Section = preamble
        image_index = 0

        def attach(level: int, heading_text: str) -> Section:
            nonlocal current
            new_section = Section(heading=heading_text, heading_level=HeadingLevel(level))
            while stack and stack[-1][0] >= level:
                stack.pop()
            if stack:
                stack[-1][1].subsections.append(new_section)
            else:
                root_sections.append(new_section)
            stack.append((level, new_section))
            current = new_section
            return new_section

        def parse_table(table: "Table") -> TableElement:
            n_rows = len(table.rows)
            n_cols = len(table.columns) if table.rows else 0
            cells: list[TableCell] = []
            for r, row in enumerate(table.rows):
                for c, cell in enumerate(row.cells):
                    cells.append(
                        TableCell(text=cell.text, row=r, col=c, is_header=(r == 0))
                    )
            return TableElement(n_rows=n_rows, n_cols=n_cols, cells=cells)

        def extract_hyperlinks(paragraph: "Paragraph") -> list[str]:
            urls = []
            for rel in paragraph.part.rels.values():
                if "hyperlink" in rel.reltype and rel.target_ref.startswith("http"):
                    urls.append(rel.target_ref)
            return urls

        def extract_images(paragraph: "Paragraph") -> list[ImageElement]:
            nonlocal image_index
            images: list[ImageElement] = []
            drawings = paragraph._p.findall(
                ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
            )
            for blip in drawings:
                embed_id = blip.get(
                    "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                )
                if not embed_id:
                    continue
                try:
                    image_part = paragraph.part.related_parts[embed_id]
                except KeyError:
                    continue
                image_index += 1
                out_dir = file_path.parent / f"{file_path.stem}_assets"
                out_dir.mkdir(parents=True, exist_ok=True)
                ext = image_part.content_type.split("/")[-1] or "png"
                out_path = out_dir / f"image_{image_index}.{ext}"
                try:
                    out_path.write_bytes(image_part.blob)
                    stored_path = str(out_path)
                except OSError:
                    stored_path = None
                images.append(ImageElement(stored_path=stored_path))
            return images

        for child in document.element.body.iterchildren():
            if isinstance(child, CT_P):
                paragraph = Paragraph(child, document)
                text = paragraph.text.strip()
                heading_level = _heading_level_from_style(paragraph.style.name if paragraph.style else "")

                images = extract_images(paragraph)
                urls = [UrlElement(url=u) for u in extract_hyperlinks(paragraph)]
                urls.extend(UrlElement(url=u) for u in _URL_RE.findall(text))

                if heading_level is not None and text:
                    section = attach(heading_level.value, text)
                    section.images.extend(images)
                    section.urls.extend(urls)
                    continue

                if not text and not images:
                    continue

                is_list = bool(paragraph.style and "list" in paragraph.style.name.lower())
                if text:
                    current.blocks.append(
                        ContentBlock(
                            type=ElementType.LIST_ITEM if is_list else ElementType.PARAGRAPH,
                            text=text,
                        )
                    )
                current.images.extend(images)
                current.urls.extend(urls)

            elif isinstance(child, CT_Tbl):
                table = Table(child, document)
                current.tables.append(parse_table(table))

        if not preamble.blocks and not preamble.subsections and not preamble.tables:
            root_sections.remove(preamble)

        core_props = document.core_properties
        metadata = DocumentMetadata(
            source_filename=file_path.name,
            file_type=SupportedFileType.DOCX,
            file_size_bytes=file_path.stat().st_size,
            author=core_props.author or None,
            title=core_props.title or None,
        )
        doc = StructuredDocument(metadata=metadata, sections=root_sections)
        logger.info(
            "Parsed DOCX '%s': %d top-level sections, %d tables, %d images",
            file_path.name,
            len(root_sections),
            doc.total_table_count(),
            doc.total_image_count(),
        )
        return doc
