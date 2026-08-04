"""
Export formatting helpers for Phase 2B's download buttons.

These functions take the already-persisted bundle written by
``app.teaching_package.persistence.save_teaching_package`` (document
metadata + KnowledgeJSON + TeachingPackage) and turn it into plain
``(heading, paragraph_lines)`` sections. Both the PDF and DOCX renderers
build from these same sections, so formatting logic lives in exactly one
place - no AI generation happens here, this is presentation-only.
"""
from __future__ import annotations

from typing import Any, Dict, List, Tuple

Section = Tuple[str, List[str]]


def _get(d: Dict[str, Any], *path, default=None):
    current: Any = d
    for key in path:
        if not isinstance(current, dict) or key not in current:
            return default
        current = current[key]
    return current if current is not None else default


def build_teaching_package_sections(bundle: Dict[str, Any]) -> List[Section]:
    """Flatten a persisted bundle into an ordered list of report sections."""
    sections: List[Section] = []

    metadata = bundle.get("document_metadata") or {}
    knowledge = bundle.get("knowledge_json") or {}
    package = bundle.get("teaching_package") or {}

    overview_lines = [
        f"Subject: {metadata.get('subject') or 'Not specified'}",
        f"Grade: {metadata.get('grade') or 'Not specified'}",
        f"Topic: {metadata.get('topic') or 'Not specified'}",
        f"Chapter: {metadata.get('chapter') or 'Not specified'}",
        f"Difficulty: {metadata.get('difficulty') or 'Not specified'}",
    ]
    sections.append(("Document Overview", overview_lines))

    los = knowledge.get("learning_objectives") or []
    if los:
        sections.append(
            ("Learning Objectives", [f"- {lo.get('text', '')}" for lo in los])
        )

    concepts = knowledge.get("concepts") or []
    if concepts:
        sections.append(
            (
                "Key Concepts",
                [f"- {c.get('name', '')}: {c.get('description', '')}" for c in concepts],
            )
        )

    lesson_plan = package.get("lesson_plan")
    if lesson_plan:
        lines = [f"Total periods: {lesson_plan.get('total_periods', '')}"]
        if lesson_plan.get("pacing_rationale"):
            lines.append(lesson_plan["pacing_rationale"])
        for period in lesson_plan.get("periods") or []:
            lines.append("")
            lines.append(
                f"Period {period.get('period_number')} "
                f"({period.get('duration_minutes')} min): {period.get('title', '')}"
            )
            if period.get("summary"):
                lines.append(period["summary"])
        sections.append(("Lesson Plan", lines))

    entry_ticket = package.get("entry_ticket")
    if entry_ticket and entry_ticket.get("items"):
        lines = []
        for item in entry_ticket["items"]:
            lines.append(f"Period {item.get('period_number')}: {item.get('question', '')}")
            if item.get("expected_answer"):
                lines.append(f"  Expected answer: {item['expected_answer']}")
        sections.append(("Entry Tickets", lines))

    teacher_script = package.get("teacher_script")
    if teacher_script and teacher_script.get("items"):
        lines = []
        for item in teacher_script["items"]:
            lines.append(f"Period {item.get('period_number')}")
            if item.get("introduction"):
                lines.append(f"  Introduction: {item['introduction']}")
            if item.get("explanation"):
                lines.append(f"  Explanation: {item['explanation']}")
            if item.get("closure"):
                lines.append(f"  Closure: {item['closure']}")
            if item.get("mentor_moment"):
                lines.append(f"  Mentor Moment: {item['mentor_moment']}")
        sections.append(("Teacher Script", lines))

    blackboard = package.get("blackboard_notes")
    if blackboard and blackboard.get("items"):
        lines = []
        for item in blackboard["items"]:
            lines.append(f"Period {item.get('period_number')}")
            for point in item.get("bullet_points") or []:
                lines.append(f"  - {point}")
        sections.append(("Blackboard Notes", lines))

    activity = package.get("classroom_activity")
    if activity and activity.get("items"):
        lines = []
        for item in activity["items"]:
            lines.append(
                f"Period {item.get('period_number')}: {item.get('title', '')} "
                f"({item.get('duration_minutes')} min, {item.get('activity_type', '')})"
            )
            if item.get("instructions"):
                lines.append(f"  Instructions: {item['instructions']}")
            if item.get("success_criteria"):
                lines.append(f"  Success criteria: {item['success_criteria']}")
        sections.append(("Classroom Activities", lines))

    assessment = package.get("assessment")
    if assessment:
        lines = []
        for i, mcq in enumerate(assessment.get("mcqs") or [], start=1):
            lines.append(f"MCQ {i}: {mcq.get('question', '')}")
            for opt in mcq.get("options") or []:
                lines.append(f"  - {opt}")
            if mcq.get("correct_option"):
                lines.append(f"  Correct: {mcq['correct_option']}")
        for i, sa in enumerate(assessment.get("short_answer") or [], start=1):
            lines.append(f"Short Answer {i}: {sa.get('question', '')}")
        for i, la in enumerate(assessment.get("long_answer") or [], start=1):
            lines.append(f"Long Answer {i}: {la.get('question', '')}")
        for i, num in enumerate(assessment.get("numerical") or [], start=1):
            lines.append(f"Numerical {i}: {num.get('question', '')}")
        if assessment.get("rubric"):
            lines.append(f"Rubric: {assessment['rubric']}")
        if lines:
            sections.append(("Assessment", lines))

    exit_ticket = package.get("exit_ticket")
    if exit_ticket and exit_ticket.get("items"):
        lines = [
            f"Period {item.get('period_number')}: {item.get('question', '')}"
            for item in exit_ticket["items"]
        ]
        sections.append(("Exit Tickets", lines))

    homework = package.get("homework")
    if homework and homework.get("items"):
        lines = []
        for item in homework["items"]:
            lines.append(f"Period {item.get('period_number')}")
            for task in item.get("tasks") or []:
                lines.append(f"  - {task}")
        sections.append(("Homework", lines))

    guidance = package.get("teacher_guidance")
    if guidance:
        lines = []
        if guidance.get("motivation_of_the_day"):
            lines.append(f"Motivation of the day: {guidance['motivation_of_the_day']}")
        if guidance.get("key_takeaway"):
            lines.append(f"Key takeaway: {guidance['key_takeaway']}")
        for tip in guidance.get("teaching_tips") or []:
            lines.append(f"Tip: {tip}")
        for misc in guidance.get("misconception_guidance") or []:
            lines.append(
                f"Misconception ({misc.get('severity', 'low')}): {misc.get('misconception', '')}"
            )
            if misc.get("remedial_action"):
                lines.append(f"  Remedial action: {misc['remedial_action']}")
        if lines:
            sections.append(("Teacher Guidance", lines))

    if package.get("generation_errors"):
        lines = [f"{name}: {msg}" for name, msg in package["generation_errors"].items()]
        sections.append(("Modules That Could Not Be Generated", lines))

    return sections


def render_docx_bytes(bundle: Dict[str, Any]) -> bytes:
    """Render a persisted bundle into a .docx Teacher Knowledge Package document."""
    import io

    from docx import Document
    from docx.shared import Pt

    metadata = bundle.get("document_metadata") or {}
    document = Document()

    title = metadata.get("topic") or metadata.get("chapter") or "Teaching Package"
    document.add_heading(title, level=0)
    subtitle = document.add_paragraph()
    subtitle.add_run(
        f"Subject: {metadata.get('subject') or 'N/A'}  |  "
        f"Grade: {metadata.get('grade') or 'N/A'}  |  "
        f"Chapter: {metadata.get('chapter') or 'N/A'}"
    ).italic = True

    for heading, lines in build_teaching_package_sections(bundle):
        document.add_heading(heading, level=1)
        for line in lines:
            paragraph = document.add_paragraph(line)
            paragraph.style.font.size = Pt(11)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def render_pdf_bytes(bundle: Dict[str, Any]) -> bytes:
    """Render a persisted bundle into a simple, paginated .pdf Teacher Knowledge
    Package document using PyMuPDF (already a parsing dependency - no new
    package is introduced for this)."""
    import textwrap

    import fitz  # PyMuPDF

    metadata = bundle.get("document_metadata") or {}
    title = metadata.get("topic") or metadata.get("chapter") or "Teaching Package"

    page_width, page_height = fitz.paper_size("a4")
    margin = 50
    line_height = 14
    max_line_width = 95

    doc = fitz.open()
    page = doc.new_page(width=page_width, height=page_height)
    y = margin

    def new_page():
        nonlocal page, y
        page = doc.new_page(width=page_width, height=page_height)
        y = margin

    def write_line(text: str, *, size: float = 11, bold: bool = False):
        nonlocal y
        if y > page_height - margin:
            new_page()
        page.insert_text(
            (margin, y),
            text,
            fontsize=size,
            fontname="helv" if not bold else "hebo",
        )
        y += line_height if size <= 12 else line_height + 8

    write_line(title, size=18, bold=True)
    write_line(
        f"Subject: {metadata.get('subject') or 'N/A'}   "
        f"Grade: {metadata.get('grade') or 'N/A'}   "
        f"Chapter: {metadata.get('chapter') or 'N/A'}",
        size=10,
    )
    y += 10

    for heading, lines in build_teaching_package_sections(bundle):
        y += 8
        write_line(heading, size=14, bold=True)
        for line in lines:
            for wrapped in textwrap.wrap(line, width=max_line_width) or [""]:
                write_line(wrapped, size=10.5)

    return doc.tobytes()
